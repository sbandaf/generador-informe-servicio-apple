from fastapi import FastAPI, UploadFile
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
import pandas as pd
from docx import Document
from docx.shared import Inches
import tempfile
from io import BytesIO
from datetime import datetime
import matplotlib.pyplot as plt

app = FastAPI()

# Permitir llamadas desde el frontend (Vite)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

MESES_ES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
}


def mes_anio_es(dt: datetime) -> str:
    return f"{MESES_ES[dt.month]} {dt.year}"


def parse_excel_datetime(s: pd.Series) -> pd.Series:
    """Convierte fechas que pueden venir como texto/fecha o como número Excel (ej: 46142.4)."""
    if pd.api.types.is_numeric_dtype(s):
        return pd.to_datetime(s, unit="D", origin="1899-12-30", errors="coerce")
    return pd.to_datetime(s, errors="coerce")


def filtrar_bch_apple(df: pd.DataFrame) -> pd.DataFrame:
    """Si existe columna 'Grupo de asignación', filtra BCH-APPLE (case-insensitive)."""
    if "Grupo de asignación" in df.columns:
        return df[df["Grupo de asignación"].astype(str).str.upper().eq("BCH-APPLE")].copy()
    return df.copy()


def serie_mensual(df: pd.DataFrame, col_fecha: str, ultimos: int = 6) -> pd.Series:
    df = df.copy()
    df[col_fecha] = parse_excel_datetime(df[col_fecha])
    df = df.dropna(subset=[col_fecha])
    df["Mes"] = df[col_fecha].dt.to_period("M").astype(str)
    s = df.groupby("Mes").size().sort_index()
    if len(s) > ultimos:
        s = s.iloc[-ultimos:]
    return s


def variacion_pct(actual: int, anterior: int):
    if anterior == 0:
        return None
    return round(((actual - anterior) / anterior) * 100, 1)


def insertar_imagen_en_placeholder(doc: Document, placeholder: str, image_path: str, width_inches: float = 5.5):
    for p in doc.paragraphs:
        if placeholder in p.text:
            p.clear()
            run = p.add_run()
            run.add_picture(image_path, width=Inches(width_inches))
            return True
    return False


def reemplazar_texto_en_doc(doc: Document, mapping: dict):
    # Párrafos
    for p in doc.paragraphs:
        for k, v in mapping.items():
            if k in p.text:
                p.text = p.text.replace(k, v)
    # Tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k, v in mapping.items():
                    if k in cell.text:
                        cell.text = cell.text.replace(k, v)


def insertar_tablas_top5_en_placeholder(doc: Document, placeholder: str, top5_por_mes: dict):
    """Reemplaza el párrafo con placeholder por tablas (una por mes)."""
    target_p = None
    for p in doc.paragraphs:
        if placeholder in p.text:
            target_p = p
            break
    if target_p is None:
        return False

    # Limpia el placeholder
    target_p.text = ""

    # Insertamos al final del documento (más robusto con python-docx) y dejamos un encabezado en el anexo.
    # Si prefieres inserción exacta, se puede hacer con XML, pero esto es estable y mantenible.
    for mes_key, df_top in top5_por_mes.items():
        doc.add_paragraph("")
        doc.add_heading(f"Mes: {mes_key}", level=2)
        t = doc.add_table(rows=1, cols=4)
        t.style = "Light Grid"
        headers = ["Fecha", "INC", "RITM", "Total"]
        for i, h in enumerate(headers):
            t.cell(0, i).text = h
        for _, r in df_top.iterrows():
            row = t.add_row().cells
            row[0].text = str(r["Fecha"])
            row[1].text = str(int(r["INC"]))
            row[2].text = str(int(r["RITM"]))
            row[3].text = str(int(r["Total"]))

    return True


def top5_dias_por_mes_bch_apple(inc: pd.DataFrame, ritm: pd.DataFrame, meses: list, top: int = 5):
    def prep(df, col_fecha, label):
        df = filtrar_bch_apple(df)
        df[col_fecha] = parse_excel_datetime(df[col_fecha])
        df = df.dropna(subset=[col_fecha])
        df["Mes"] = df[col_fecha].dt.to_period("M").astype(str)
        df["Fecha"] = df[col_fecha].dt.date
        daily = df.groupby(["Mes", "Fecha"]).size().reset_index(name=label)
        return daily

    inc_d = prep(inc, "Creado", "INC")
    ritm_d = prep(ritm, "Creado", "RITM")

    merged = pd.merge(inc_d, ritm_d, on=["Mes", "Fecha"], how="outer").fillna(0)
    merged["INC"] = merged["INC"].astype(int)
    merged["RITM"] = merged["RITM"].astype(int)
    merged["Total"] = merged["INC"] + merged["RITM"]

    out = {}
    for m in meses:
        sub = merged[merged["Mes"] == m].sort_values("Total", ascending=False).head(top).copy()
        sub["Fecha"] = pd.to_datetime(sub["Fecha"]).dt.strftime("%d-%m-%Y")
        out[m] = sub[["Fecha", "INC", "RITM", "Total"]]
    return out


@app.get("/")
def root():
    return {"status": "ok", "message": "API Generador Informe Servicio Apple"}


@app.post("/generar-informe")
async def generar_informe(
    INC: UploadFile,
    RITM: UploadFile,
    INC_Abiertos: UploadFile,
    RITM_Abiertos: UploadFile,
):
    # Leer archivos
    inc = pd.read_excel(BytesIO(await INC.read()))
    ritm = pd.read_excel(BytesIO(await RITM.read()))
    inc_ab = pd.read_excel(BytesIO(await INC_Abiertos.read()))
    ritm_ab = pd.read_excel(BytesIO(await RITM_Abiertos.read()))

    # Filtro servicio (si existe columna)
    inc = filtrar_bch_apple(inc)
    ritm = filtrar_bch_apple(ritm)
    inc_ab = filtrar_bch_apple(inc_ab)
    ritm_ab = filtrar_bch_apple(ritm_ab)

    # Fechas y tiempos
    inc["Creado"] = parse_excel_datetime(inc["Creado"]) if "Creado" in inc.columns else pd.NaT
    inc["Cerrado"] = parse_excel_datetime(inc["Cerrado"]) if "Cerrado" in inc.columns else pd.NaT
    inc["Tiempo"] = (inc["Cerrado"] - inc["Creado"]).dt.days

    ritm["Creado"] = parse_excel_datetime(ritm["Creado"]) if "Creado" in ritm.columns else pd.NaT
    cierre_col = "Cerrado" if "Cerrado" in ritm.columns else ("Actualizado" if "Actualizado" in ritm.columns else None)
    ritm["Cerrado_calc"] = parse_excel_datetime(ritm[cierre_col]) if cierre_col else pd.NaT
    ritm["Tiempo"] = (ritm["Cerrado_calc"] - ritm["Creado"]).dt.days

    # Series 6 meses
    inc_m = serie_mensual(inc, "Creado", ultimos=6) if "Creado" in inc.columns else pd.Series(dtype=int)
    ritm_m = serie_mensual(ritm, "Creado", ultimos=6) if "Creado" in ritm.columns else pd.Series(dtype=int)
    meses_union = sorted(set(inc_m.index).union(set(ritm_m.index)))

    # Periodo del informe = último mes con datos
    if meses_union:
        periodo_key = meses_union[-1]  # "YYYY-MM"
        periodo_dt = pd.Period(periodo_key).to_timestamp()
        periodo_str = mes_anio_es(periodo_dt.to_pydatetime())
    else:
        now = datetime.now()
        periodo_key = now.strftime("%Y-%m")
        periodo_str = mes_anio_es(now)

    # Datos del periodo
    inc_periodo = inc[inc["Creado"].dt.to_period("M").astype(str) == periodo_key] if "Creado" in inc.columns else inc
    ritm_periodo = ritm[ritm["Creado"].dt.to_period("M").astype(str) == periodo_key] if "Creado" in ritm.columns else ritm

    inc_count = int(len(inc_periodo))
    ritm_count = int(len(ritm_periodo))

    # Variación MoM (total y por tipo)
    if len(meses_union) >= 2:
        prev_key = meses_union[-2]
        inc_var = variacion_pct(int(inc_m.get(periodo_key, 0)), int(inc_m.get(prev_key, 0)))
        ritm_var = variacion_pct(int(ritm_m.get(periodo_key, 0)), int(ritm_m.get(prev_key, 0)))
        total_var = variacion_pct(int(inc_m.get(periodo_key, 0)) + int(ritm_m.get(periodo_key, 0)),
                                 int(inc_m.get(prev_key, 0)) + int(ritm_m.get(prev_key, 0)))
    else:
        inc_var = ritm_var = total_var = None

    # Top 5 días por mes (últimos 6)
    top5_por_mes = top5_dias_por_mes_bch_apple(inc, ritm, meses_union[-6:] if meses_union else [periodo_key], top=5)

    # Gráfico periodo
    plt.figure(figsize=(6, 4))
    plt.bar(["INC", "RITM"], [inc_count, ritm_count], color=["#E60012", "#6B7280"])
    plt.title(f"Carga del periodo ({periodo_str})")
    plt.ylabel("Cantidad de tickets")
    plt.tight_layout()
    img_bar = os.path.join(tempfile.gettempdir(), "grafico_periodo.png")
    plt.savefig(img_bar)
    plt.close()

    # Gráfico tendencia 6m
    df_trend = pd.DataFrame({
        "Mes": meses_union,
        "INC": [int(inc_m.get(m, 0)) for m in meses_union],
        "RITM": [int(ritm_m.get(m, 0)) for m in meses_union],
    })
    plt.figure(figsize=(7, 3.8))
    if not df_trend.empty:
        plt.plot(df_trend["Mes"], df_trend["INC"], marker="o", label="INC", color="#E60012")
        plt.plot(df_trend["Mes"], df_trend["RITM"], marker="o", label="RITM", color="#6B7280")
        plt.xticks(rotation=30, ha="right")
        plt.title("Tendencia de carga (últimos 6 meses)")
        plt.ylabel("Cantidad de tickets")
        plt.legend()
        plt.tight_layout()
    img_trend = os.path.join(tempfile.gettempdir(), "tendencia_6m.png")
    plt.savefig(img_trend)
    plt.close()

    # Conclusión ejecutiva (larga, técnica y orientada a gerencia)
    total_periodo = inc_count + ritm_count
    abiertos = len(inc_ab) + len(ritm_ab)

    def fmt_var(v):
        if v is None:
            return "s/d"
        sign = "+" if v > 0 else ""
        return f"{sign}{v}%"

    mes_mayor = None
    if meses_union:
        totales_m = {m: int(inc_m.get(m, 0)) + int(ritm_m.get(m, 0)) for m in meses_union}
        mes_mayor = max(totales_m, key=totales_m.get)
        pico = totales_m[mes_mayor]
    else:
        pico = total_periodo

    # identificar picos diarios del periodo
    top_periodo = top5_por_mes.get(periodo_key)
    if top_periodo is not None and not top_periodo.empty:
        pico_dia = top_periodo.iloc[0].to_dict()
        pico_txt = f"El pico diario del periodo se observó el {pico_dia['Fecha']} con {pico_dia['Total']} tickets (INC={pico_dia['INC']}, RITM={pico_dia['RITM']})."
    else:
        pico_txt = "No fue posible identificar picos diarios del periodo con la información disponible."

    conclusion = (
        f"Resumen ejecutivo del periodo {periodo_str}: se registraron {total_periodo} tickets en total (INC={inc_count}, RITM={ritm_count}). "
        f"La carga mensual total presenta una variación de {fmt_var(total_var)} respecto del mes anterior; "
        f"para INC la variación fue {fmt_var(inc_var)} y para RITM {fmt_var(ritm_var)}.\n\n"
        f"Desde la perspectiva operativa, el volumen del periodo se concentra en picos diarios específicos (Top 5 por mes en el anexo), "
        f"lo que sugiere demanda no homogénea y necesidad de planificación por ventanas de alta carga. {pico_txt}\n\n"
        f"En eficiencia de atención, el tiempo promedio de resolución fue de {round(inc_periodo['Tiempo'].mean(), 1)} días para INC "
        f"y {round(ritm_periodo['Tiempo'].mean(), 1)} días para RITM (según hito de cierre disponible). "
        f"Se recomienda vigilar outliers que eleven el promedio y reforzar acciones preventivas (automatización/estandarización) en los casos recurrentes.\n\n"
        f"Backlog: al cierre del periodo se registran {abiertos} tickets abiertos (INC={len(inc_ab)}, RITM={len(ritm_ab)}). "
        f"Si la tendencia de carga se mantiene al alza, se recomienda revisar capacidad (dotación/turnos) y priorización, "
        f"apoyándose en el análisis de picos diarios y en la evolución de los últimos 6 meses."
    )

    # Mapping placeholders
    now = datetime.now()
    mapping = {
        "{{MES_ANIO}}": periodo_str,
        "{{FECHA_EMISION}}": now.strftime("%d-%m-%Y"),
        "{{TOTAL_INC}}": str(inc_count),
        "{{TOTAL_RITM}}": str(ritm_count),
        "{{INC_ABIERTOS}}": str(len(inc_ab)),
        "{{RITM_ABIERTOS}}": str(len(ritm_ab)),
        "{{INC_MES}}": str(inc_count),
        "{{RITM_MES}}": str(ritm_count),
        "{{TIEMPO_INC}}": str(round(inc_periodo['Tiempo'].mean(), 1)),
        "{{TIEMPO_RITM}}": str(round(ritm_periodo['Tiempo'].mean(), 1)),
        "{{CONCLUSION_DETALLADA}}": conclusion,
        "{{TABLAS_TOP5}}": "",  # se reemplaza con tablas
    }

    # Crear documento desde plantilla Ricoh
    doc = Document("Plantilla_Informe_Mensual_Servicio_Apple_RICOH.docx")
    reemplazar_texto_en_doc(doc, mapping)

    # Insertar gráficos en placeholders
    insertar_imagen_en_placeholder(doc, "ESPACIO PARA GRÁFICO", img_bar, width_inches=5.5)
    insertar_imagen_en_placeholder(doc, "ESPACIO PARA TENDENCIA 6M", img_trend, width_inches=6.5)

    # Insertar tablas Top5 (en el anexo)
    insertar_tablas_top5_en_placeholder(doc, "{{TABLAS_TOP5}}", top5_por_mes)

    # Nombre de archivo con mes/año
    safe_period = periodo_str.replace(" ", "_")
    filename = f"Informe_Servicio_Apple_{safe_period}.docx"

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)

    return FileResponse(tmp.name, filename=filename)
